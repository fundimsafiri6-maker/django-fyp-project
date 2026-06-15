from django.shortcuts import render, redirect
from django.contrib.auth import logout, authenticate, login
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.db.models import Count, Q
from django.db import transaction
from django.core.mail import send_mail, EmailMultiAlternatives
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.conf import settings
from django.db import IntegrityError
import requests
from bs4 import BeautifulSoup
import json
import re
import logging
from .models import User, EmailVerificationToken
from .decorators import admin_required, staff_or_admin_required
from complaints.models import Complaint
from django.views.decorators.csrf import csrf_exempt
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

@csrf_exempt
def login_view(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")

        # Try to find user by username
        try:
            user = User.objects.get(username=username)
        except User.DoesNotExist:
            messages.error(request, "Invalid credentials")
            return render(request, "accounts/login.html")

        # Check if student's email is not verified
        if user.role == 'student' and not user.is_email_verified:
            messages.error(
                request,
                f"Please verify your email before signing in. Check your inbox for the verification link sent to {user.email}. If you didn't receive it, you can resend the verification email."
            )
            return render(request, "accounts/login.html", {"show_resend": True, "unverified_email": user.email})
        
        # Now authenticate the user with password check
        authenticated_user = authenticate(request, username=username, password=password)
        if authenticated_user is not None:
            login(request, authenticated_user)
            full_name = authenticated_user.get_full_name() if authenticated_user.get_full_name() else authenticated_user.username
            messages.success(request, f"Welcome back, {full_name}!")
            # Redirect based on user role
            if authenticated_user.role == 'admin':
                return redirect('admin_dashboard')
            elif authenticated_user.role == 'staff':
                return redirect('staff_dashboard')
            else:  # student
                return redirect('student_dashboard')
        else:
            messages.error(request, "Invalid credentials")
    
    return render(request, "accounts/login.html")

def logout_view(request):
    logout(request)
    return redirect('login')


# ================= DASHBOARDS =================

@login_required(login_url='login')
def admin_dashboard(request):
    """Admin dashboard with system-wide statistics"""
    # Get all complaints
    all_complaints = Complaint.objects.all()
    
    # Calculate statistics
    total_complaints = all_complaints.count()
    pending = all_complaints.filter(status='Pending').count()
    in_progress = all_complaints.filter(status='In Progress').count()
    resolved = all_complaints.filter(status='Resolved').count()
    rejected = all_complaints.filter(status='Rejected').count()
    
    # Calculate percentages for status bar
    total_status = pending + in_progress + resolved + rejected
    if total_status > 0:
        pending_percent = int((pending / total_status) * 100)
        in_progress_percent = int((in_progress / total_status) * 100)
        resolved_percent = int((resolved / total_status) * 100)
        rejected_percent = max(0, 100 - (pending_percent + in_progress_percent + resolved_percent))
    else:
        pending_percent = in_progress_percent = resolved_percent = rejected_percent = 0
    
    context = {
        'total_complaints': total_complaints,
        'pending': pending,
        'in_progress': in_progress,
        'resolved': resolved,
        'rejected': rejected,
        'pending_percent': pending_percent,
        'in_progress_percent': in_progress_percent,
        'resolved_percent': resolved_percent,
        'rejected_percent': rejected_percent,
    }
    
    return render(request, "dashboard/admin_dashboard.html", context)

@login_required(login_url='login')
def staff_dashboard(request):
    """Staff dashboard showing complaints for their department with accurate statistics"""
    if request.user.role != 'staff':
        messages.error(request, 'Access denied. Staff only.')
        return redirect('student_dashboard')
    
    # Get staff's department
    staff_department = request.user.department
    
    # Get complaints for this department
    if staff_department:
        department_complaints = Complaint.objects.filter(
            department=staff_department
        ).order_by('-created_at')
    else:
        # If staff has no department assigned, show all complaints
        department_complaints = Complaint.objects.all().order_by('-created_at')
    
    # Calculate analytics - matching EXACTLY with complaint statuses
    total_complaints = department_complaints.count()
    assigned_count = department_complaints.filter(assigned_to=request.user).count()
    pending = department_complaints.filter(status='Pending').count()
    in_progress = department_complaints.filter(status='In Progress').count()
    resolved = department_complaints.filter(status='Resolved').count()
    rejected = department_complaints.filter(status='Rejected').count()
    high_priority = department_complaints.filter(priority__in=['High', 'Urgent']).count()
    
    # Get status distribution for status bar
    total_status_complaints = pending + in_progress + resolved + rejected
    
    # Calculate percentages for status bar (prevent division by zero)
    if total_status_complaints > 0:
        pending_percent = int((pending / total_status_complaints) * 100)
        in_progress_percent = int((in_progress / total_status_complaints) * 100)
        resolved_percent = int((resolved / total_status_complaints) * 100)
        rejected_percent = max(0, 100 - (pending_percent + in_progress_percent + resolved_percent))
    else:
        pending_percent = in_progress_percent = resolved_percent = rejected_percent = 0
    
    # Get status distribution for charts
    status_distribution = {
        'Pending': pending,
        'In Progress': in_progress,
        'Resolved': resolved,
        'Rejected': rejected,
    }
    
    # Get recent complaints (unfiltered by staff)
    recent_complaints = department_complaints[:5]
    
    context = {
        'department': staff_department,
        'department_display': dict(Complaint.DEPARTMENT_CHOICES).get(staff_department, 'Unassigned'),
        'total_complaints': total_complaints,
        'assigned_count': assigned_count,
        'pending': pending,
        'in_progress': in_progress,
        'resolved': resolved,
        'rejected': rejected,
        'high_priority': high_priority,
        'status_distribution': status_distribution,
        'recent_complaints': recent_complaints,
        # Percentages for status bar
        'pending_percent': pending_percent,
        'in_progress_percent': in_progress_percent,
        'resolved_percent': resolved_percent,
        'rejected_percent': rejected_percent,
    }
    return render(request, "dashboard/staff_dashboard.html", context)

@login_required(login_url='login')
def hod_dashboard(request):
    """HOD dashboard showing HOD-classified complaints with statistics"""
    if request.user.role != 'hod':
        messages.error(request, 'Access denied. HOD only.')
        return redirect('student_dashboard')
    
    # Get HOD-classified complaints
    hod_complaints = Complaint.objects.filter(
        department='hod'
    ).order_by('-created_at')
    
    # Calculate analytics
    total_complaints = hod_complaints.count()
    assigned_count = hod_complaints.filter(assigned_to=request.user).count()
    pending = hod_complaints.filter(status='Pending').count()
    in_progress = hod_complaints.filter(status='In Progress').count()
    resolved = hod_complaints.filter(status='Resolved').count()
    rejected = hod_complaints.filter(status='Rejected').count()
    high_priority = hod_complaints.filter(priority__in=['High', 'Urgent']).count()
    
    # Get priority breakdown for charts
    urgent = hod_complaints.filter(priority='Urgent').count()
    high = hod_complaints.filter(priority='High').count()
    medium = hod_complaints.filter(priority='Medium').count()
    low = hod_complaints.filter(priority='Low').count()
    
    # Get status distribution for status bar
    total_status_complaints = pending + in_progress + resolved + rejected
    
    # Calculate percentages for status bar (prevent division by zero)
    if total_status_complaints > 0:
        pending_percent = int((pending / total_status_complaints) * 100)
        in_progress_percent = int((in_progress / total_status_complaints) * 100)
        resolved_percent = int((resolved / total_status_complaints) * 100)
        rejected_percent = max(0, 100 - (pending_percent + in_progress_percent + resolved_percent))
    else:
        pending_percent = in_progress_percent = resolved_percent = rejected_percent = 0
    
    # Get status distribution for charts
    status_distribution = {
        'Pending': pending,
        'In Progress': in_progress,
        'Resolved': resolved,
        'Rejected': rejected,
    }
    
    # Get recent complaints
    recent_complaints = hod_complaints[:5]
    
    context = {
        'department': 'hod',
        'department_display': 'HOD (Head of Department)',
        'total_complaints': total_complaints,
        'assigned_count': assigned_count,
        'pending': pending,
        'in_progress': in_progress,
        'resolved': resolved,
        'rejected': rejected,
        'high_priority': high_priority,
        'urgent': urgent,
        'high': high,
        'medium': medium,
        'low': low,
        'status_distribution': status_distribution,
        'recent_complaints': recent_complaints,
        # Percentages for status bar
        'pending_percent': pending_percent,
        'in_progress_percent': in_progress_percent,
        'resolved_percent': resolved_percent,
        'rejected_percent': rejected_percent,
    }
    return render(request, "dashboard/hod_dashboard.html", context)

@login_required(login_url='login')
def academic_department_dashboard(request):
    """Academic Department staff dashboard"""
    if request.user.role != 'staff' or request.user.department != 'academic':
        messages.error(request, 'Access denied. Academic department staff only.')
        return redirect('staff_dashboard')

    return _get_department_dashboard(request, 'academic', 'Academic Department', 'dashboard/academic_dashboard.html')

@login_required(login_url='login')
def ict_department_dashboard(request):
    """ICT Department staff dashboard"""
    if request.user.role != 'staff' or request.user.department != 'ict':
        messages.error(request, 'Access denied. ICT department staff only.')
        return redirect('staff_dashboard')

    return _get_department_dashboard(request, 'ict', 'ICT Department', 'dashboard/ict_dashboard.html')

def _get_department_dashboard(request, department_code, department_name, template):
    """
    Generic department dashboard view logic
    
    Args:
        request: HTTP request
        department_code: Department code (academic, ict, other)
        department_name: Display name for department
        template: Template path to render
    """
    # Get all complaints for this department
    department_complaints = Complaint.objects.filter(
        department=department_code
    ).order_by('-created_at')
    
    # Calculate analytics
    assigned_to_me = department_complaints.filter(assigned_to=request.user)
    pending = department_complaints.filter(status='Pending')
    in_progress = department_complaints.filter(status='In Progress')
    resolved = department_complaints.filter(status='Resolved')
    rejected = department_complaints.filter(status='Rejected')
    high_priority = department_complaints.filter(priority__in=['High', 'Urgent'])
    urgent = department_complaints.filter(priority='Urgent')
    ai_classified = department_complaints.filter(ai_classified=True)
    
    # Get all staff in this department
    department_staff = User.objects.filter(
        role='staff',
        department=department_code
    ).annotate(
        complaint_count=Count('assigned_complaints')
    ).order_by('-complaint_count')
    
    # Get unassigned complaints
    unassigned_complaints = department_complaints.filter(assigned_to__isnull=True)
    
    # Calculate percentages for status bar
    pending_count = pending.count()
    in_progress_count = in_progress.count()
    resolved_count = resolved.count()
    rejected_count = rejected.count()
    total_status_complaints = pending_count + in_progress_count + resolved_count + rejected_count
    
    if total_status_complaints > 0:
        pending_percent = int((pending_count / total_status_complaints) * 100)
        in_progress_percent = int((in_progress_count / total_status_complaints) * 100)
        resolved_percent = int((resolved_count / total_status_complaints) * 100)
        rejected_percent = max(0, 100 - (pending_percent + in_progress_percent + resolved_percent))
    else:
        pending_percent = in_progress_percent = resolved_percent = rejected_percent = 0
    
    # Prepare context
    context = {
        'department_code': department_code,
        'department_name': department_name,
        'total_complaints': department_complaints.count(),
        'assigned_to_me_count': assigned_to_me.count(),
        'pending_count': pending_count,
        'in_progress_count': in_progress_count,
        'resolved_count': resolved_count,
        'rejected_count': rejected_count,
        'high_priority_count': high_priority.count(),
        'urgent_count': urgent.count(),
        'ai_classified_count': ai_classified.count(),
        'unassigned_count': unassigned_complaints.count(),
        
        # Percentages for status bar
        'pending_percent': pending_percent,
        'in_progress_percent': in_progress_percent,
        'resolved_percent': resolved_percent,
        'rejected_percent': rejected_percent,
        
        # Recent items
        'recent_assigned': assigned_to_me.order_by('-updated_at')[:5],
        'recent_pending': pending.order_by('-created_at')[:5],
        'unassigned_complaints': unassigned_complaints.order_by('-created_at')[:10],
        'urgent_complaints': urgent.order_by('-created_at')[:5],
        
        # Statistics
        'department_staff': department_staff,
        'status_distribution': {
            'Pending': pending_count,
            'In Progress': in_progress_count,
            'Resolved': resolved_count,
            'Rejected': rejected_count,
        },
        'priority_distribution': {
            'Low': department_complaints.filter(priority='Low').count(),
            'Medium': department_complaints.filter(priority='Medium').count(),
            'High': department_complaints.filter(priority='High').count(),
            'Urgent': urgent.count(),
        },
    }
    
    return render(request, template, context)

@login_required(login_url='login')
def student_dashboard(request):
    """Student dashboard showing their own complaint statistics"""
    # Get complaints submitted by this student
    my_complaints = Complaint.objects.filter(user=request.user).order_by('-created_at')
    
    # Calculate statistics
    total_complaints = my_complaints.count()
    pending = my_complaints.filter(status='Pending').count()
    in_progress = my_complaints.filter(status='In Progress').count()
    resolved = my_complaints.filter(status='Resolved').count()
    rejected = my_complaints.filter(status='Rejected').count()
    
    # Calculate percentages for status bar
    total_status = pending + in_progress + resolved + rejected
    if total_status > 0:
        pending_percent = int((pending / total_status) * 100)
        in_progress_percent = int((in_progress / total_status) * 100)
        resolved_percent = int((resolved / total_status) * 100)
        rejected_percent = max(0, 100 - (pending_percent + in_progress_percent + resolved_percent))
    else:
        pending_percent = in_progress_percent = resolved_percent = rejected_percent = 0
    
    context = {
        'total_complaints': total_complaints,
        'pending': pending,
        'in_progress': in_progress,
        'resolved': resolved,
        'rejected': rejected,
        'pending_percent': pending_percent,
        'in_progress_percent': in_progress_percent,
        'resolved_percent': resolved_percent,
        'rejected_percent': rejected_percent,
        'my_complaints': my_complaints[:5],  # Recent complaints
    }
    
    return render(request, "dashboard/student_dashboard.html", context)


# ================= EXTRA PAGES =================

def assign_complaint(request):
    return render(request, "dashboard/assign_complaints.html")

def statistics_page(request):
    return render(request, "dashboard/complaints_stats.html")

def report_page(request):
    """Generate and display system report as PDF or Word document"""
    format_type = request.GET.get('format', 'pdf')  # Default to PDF
    
    # Get complaint statistics
    all_complaints = Complaint.objects.all()
    total_complaints = all_complaints.count()
    pending_complaints = all_complaints.filter(status='Pending').count()
    in_progress_complaints = all_complaints.filter(status='In Progress').count()
    resolved_complaints = all_complaints.filter(status='Resolved').count()
    rejected_complaints = all_complaints.filter(status='Rejected').count()
    
    # Get top category
    from django.db.models import Count
    category_stats = all_complaints.values('category').annotate(count=Count('id')).order_by('-count').first()
    top_category = category_stats['category'] if category_stats else 'N/A'
    
    # Get most active department
    dept_stats = all_complaints.values('department').annotate(count=Count('id')).order_by('-count').first()
    most_active_dept = dict(Complaint.DEPARTMENT_CHOICES).get(dept_stats['department'], 'N/A') if dept_stats else 'N/A'
    
    # Calculate average resolution time (in days)
    resolved_with_dates = all_complaints.filter(status='Resolved', updated_at__isnull=False)
    if resolved_with_dates.exists():
        from django.db.models import F, ExpressionWrapper, DurationField
        from datetime import timedelta
        resolution_times = []
        for complaint in resolved_with_dates:
            if complaint.created_at and complaint.updated_at:
                diff = complaint.updated_at - complaint.created_at
                resolution_times.append(diff.days)
        avg_resolution_time = sum(resolution_times) / len(resolution_times) if resolution_times else 0
    else:
        avg_resolution_time = 0
    
    if format_type == 'word':
        # Generate Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('System Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('Monthly Complaint Summary Report - UDOM CIVE')
        doc.add_paragraph(f'Generated: {all_complaints.first().created_at.strftime("%B %Y") if all_complaints.exists() else "N/A"}')
        doc.add_paragraph()  # Empty line
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph('This report provides a comprehensive overview of the complaint management system.')
        doc.add_paragraph()
        
        # Statistics
        doc.add_heading('Complaint Statistics Overview', level=1)
        
        stats_table = doc.add_table(rows=7, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        stats_data = [
            ['Metric', 'Value'],
            ['Total Complaints', str(total_complaints)],
            ['Pending Review', str(pending_complaints)],
            ['In Progress', str(in_progress_complaints)],
            ['Resolved', str(resolved_complaints)],
            ['Rejected', str(rejected_complaints)],
            ['Average Resolution Time', f"{avg_resolution_time:.1f} Days"],
        ]
        
        for i, row_data in enumerate(stats_data):
            row = stats_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i == 0:  # Header row
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Key Insights
        doc.add_heading('Key Insights', level=1)
        
        insights_table = doc.add_table(rows=4, cols=2)
        insights_table.style = 'Light Grid Accent 1'
        insights_data = [
            ['Insight', 'Value'],
            ['Top Category', top_category],
            ['Most Active Department', most_active_dept],
            ['Average Resolution Time', f"{avg_resolution_time:.1f} Days"],
        ]
        
        for i, row_data in enumerate(insights_data):
            row = insights_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i == 0:  # Header row
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Status Distribution
        doc.add_heading('Status Distribution', level=1)
        
        if total_complaints > 0:
            status_data = [
                ['Status', 'Count', 'Percentage'],
                ['Pending', str(pending_complaints), f"{(pending_complaints/total_complaints*100):.2f}%"],
                ['In Progress', str(in_progress_complaints), f"{(in_progress_complaints/total_complaints*100):.2f}%"],
                ['Resolved', str(resolved_complaints), f"{(resolved_complaints/total_complaints*100):.2f}%"],
                ['Rejected', str(rejected_complaints), f"{(rejected_complaints/total_complaints*100):.2f}%"],
            ]
        else:
            status_data = [
                ['Status', 'Count', 'Percentage'],
                ['Pending', str(pending_complaints), '0.00%'],
                ['In Progress', str(in_progress_complaints), '0.00%'],
                ['Resolved', str(resolved_complaints), '0.00%'],
                ['Rejected', str(rejected_complaints), '0.00%'],
            ]
        
        status_table = doc.add_table(rows=len(status_data), cols=3)
        status_table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(status_data):
            row = status_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i == 0:  # Header row
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Footer
        doc.add_paragraph('AI Complaints Management System - University of Dodoma')
        from datetime import datetime
        doc.add_paragraph(f'Report generated on {datetime.now().strftime("%B %d, %Y %H:%M")}')
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="system_report.docx"'
        
        return response
    else:
        # Generate PDF using ReportLab
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=18,
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=12
        )
        
        # Title
        story.append(Paragraph("System Report", title_style))
        story.append(Paragraph("Monthly Complaint Summary Report - UDOM CIVE", styles['Normal']))
        story.append(Spacer(1, 0.2 * inch))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        story.append(Paragraph(
            f"This report provides a comprehensive overview of the complaint management system.",
            styles['Normal']
        ))
        story.append(Spacer(1, 0.2 * inch))
        
        # Statistics table
        story.append(Paragraph("Complaint Statistics Overview", heading_style))
        
        data = [
            ['Metric', 'Value'],
            ['Total Complaints', str(total_complaints)],
            ['Pending Review', str(pending_complaints)],
            ['In Progress', str(in_progress_complaints)],
            ['Resolved', str(resolved_complaints)],
            ['Rejected', str(rejected_complaints)],
            ['Average Resolution Time', f"{avg_resolution_time:.1f} Days"],
        ]
        
        table = Table(data, colWidths=[3 * inch, 2 * inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#1e3a8a')),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.3 * inch))
        
        # Key Insights
        story.append(Paragraph("Key Insights", heading_style))
        
        insights_data = [
            ['Insight', 'Value'],
            ['Top Category', top_category],
            ['Most Active Department', most_active_dept],
            ['Average Resolution Time', f"{avg_resolution_time:.1f} Days"],
        ]
        
        insights_table = Table(insights_data, colWidths=[3 * inch, 2 * inch])
        insights_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#1e3a8a')),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(insights_table)
        story.append(Spacer(1, 0.3 * inch))
        
        # Status Distribution
        story.append(Paragraph("Status Distribution", heading_style))
        
        if total_complaints > 0:
            status_data = [
                ['Status', 'Count', 'Percentage'],
                ['Pending', str(pending_complaints), f"{(pending_complaints/total_complaints*100):.2f}%"],
                ['In Progress', str(in_progress_complaints), f"{(in_progress_complaints/total_complaints*100):.2f}%"],
                ['Resolved', str(resolved_complaints), f"{(resolved_complaints/total_complaints*100):.2f}%"],
                ['Rejected', str(rejected_complaints), f"{(rejected_complaints/total_complaints*100):.2f}%"],
            ]
        else:
            status_data = [
                ['Status', 'Count', 'Percentage'],
                ['Pending', str(pending_complaints), '0.00%'],
                ['In Progress', str(in_progress_complaints), '0.00%'],
                ['Resolved', str(resolved_complaints), '0.00%'],
                ['Rejected', str(rejected_complaints), '0.00%'],
            ]
        
        status_table = Table(status_data, colWidths=[2 * inch, 1.5 * inch, 1.5 * inch])
        status_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (2, 0), colors.HexColor('#1e3a8a')),
            ('TEXTCOLOR', (0, 0), (2, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(status_table)
        story.append(Spacer(1, 0.5 * inch))
        
        # Footer
        story.append(Paragraph("AI Complaints Management System - University of Dodoma", styles['Normal']))
        from datetime import datetime
        story.append(Paragraph(f"Report generated on {datetime.now().strftime('%B %d, %Y %H:%M')}", styles['Normal']))
        
        # Build PDF
        doc.build(story)
        
        # Create response
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="system_report.pdf"'
        
        return response

@login_required(login_url='login')
def user_profile(request):
    """Display and edit user profile"""
    if request.method == 'POST':
        user = request.user
        user.first_name = request.POST.get('first_name', user.first_name)
        user.last_name = request.POST.get('last_name', user.last_name)
        user.email = request.POST.get('email', user.email)
        user.save()
        messages.success(request, 'Profile updated successfully!')
        return redirect('user_profile')
    
    return render(request, 'accounts/profile.html', {'user': request.user})

@login_required(login_url='login')
def settings_page(request):
    """User settings page"""
    return render(request, 'accounts/settings.html')

@login_required(login_url='login')
def help_page(request):
    """Help and support page"""
    return render(request, 'accounts/help.html')

def register_view(request):
    import re
    import logging
    from django.db import transaction
    
    logger = logging.getLogger(__name__)

    
    if request.method == "POST":
        # Get and validate input
        username = request.POST.get("username", "").strip()
        email = request.POST.get("email", "").strip().lower()  # Normalize email to lowercase
        password = request.POST.get("password", "").strip()
        password_confirm = request.POST.get("password_confirm", "").strip()
        role = "student"  # Default role for security
        registration_number = request.POST.get("registration_number", "").strip()
        department = None  # No department for students
        
        # Validate inputs
        errors = []
        
        # Username validation
        if not username:
            errors.append("Username is required")
        elif len(username) < 3:
            errors.append("Username must be at least 3 characters long")
        elif not re.match(r'^[a-zA-Z0-9_]+$', username):
            errors.append("Username can only contain letters, numbers, and underscores")
        elif User.objects.filter(username=username).exists():
            errors.append("Username already exists")
        
        # Email validation
        if not email:
            errors.append("Email is required")
        elif not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email):
            errors.append("Invalid email format")
        elif User.objects.filter(email__iexact=email).exists():  # Case-insensitive check
            errors.append("Email already registered")
        
        # Password validation
        if not password:
            errors.append("Password is required")
        elif len(password) < 8:
            errors.append("Password must be at least 8 characters long")
        elif not re.search(r'[A-Z]', password):
            errors.append("Password must contain at least one uppercase letter")
        elif not re.search(r'[a-z]', password):
            errors.append("Password must contain at least one lowercase letter")
        elif not re.search(r'[0-9]', password):
            errors.append("Password must contain at least one number")
        
        # Password confirmation
        if password != password_confirm:
            errors.append("Passwords do not match")
        

        
        # Display errors if any
        if errors:
            context = {
                'error_list': errors,
                'username': username,
                'email': email,
                'registration_number': registration_number,
            }
            return render(request, "accounts/register.html", context)
        
        # Create user with transaction to ensure atomicity
        try:
            with transaction.atomic():
                user = User.objects.create_user(
                    username=username,
                    email=email,
                    password=password,
                    role=role,
                    registration_number=registration_number if registration_number else None,
                    is_active=False  # Require email verification
                )
                
                # Create verification token and send email
                try:
                    # Delete any existing verification tokens for this user
                    EmailVerificationToken.objects.filter(user=user).delete()

                    # Create new verification token
                    verification_token = EmailVerificationToken.objects.create(user=user)

                    # Send verification email
                    send_verification_email(request, user, verification_token)

                    messages.success(
                        request,
                        f"Account created successfully! We've sent a verification link to {email}. Please check your inbox and click the link to verify your email and activate your account."
                    )
                except Exception as e:
                    logger.error(f"Email sending failed for user {username}: {str(e)}", exc_info=True)
                    # If email sending fails, delete the user and show error
                    user.delete()
                    messages.error(request, f"Failed to send verification email. Please try again later.")
                    return render(request, "accounts/register.html")
                
                logger.info(f"✅ New user registered: {username} ({role})")
                return redirect('login')
        
        except Exception as e:
            logger.error(f"Error creating user {username}: {str(e)}", exc_info=True)
            messages.error(request, f"An error occurred during registration. Please try again.")
            return render(request, "accounts/register.html")
    
    return render(request, "accounts/register.html")


def send_verification_email(request, user, verification_token):
    """Send email verification link to user with robust error handling"""
    from django.core.mail import EmailMultiAlternatives
    from django.core.exceptions import ImproperlyConfigured
    import logging
    
    logger = logging.getLogger(__name__)

    # Verify email configuration
    email_host = settings.EMAIL_HOST
    email_user = settings.EMAIL_HOST_USER
    email_password = settings.EMAIL_HOST_PASSWORD
    
    if not email_host or not email_user or not email_password:
        logger.error("Email configuration incomplete: Missing EMAIL_HOST, EMAIL_HOST_USER, or EMAIL_HOST_PASSWORD")
        raise Exception("Email service is not configured. Please contact the administrator.")

    verification_url = request.build_absolute_uri(
        f'/accounts/verify-email/{verification_token.token}/'
    )

    context = {
        'user': user,
        'verification_url': verification_url,
        'username': user.username,
        'token_expires': verification_token.expires_at,
    }

    # Render HTML email
    try:
        html_message = render_to_string('accounts/email_verification.html', context)
    except Exception as e:
        logger.error(f"Failed to render HTML email template: {str(e)}")
        # Fallback to plain text only
        html_message = None

    # Plain text version from text template
    try:
        plain_message = render_to_string('accounts/email_verification.txt', context)
    except Exception as e:
        logger.warning(f"Failed to render text email template: {str(e)}")
        plain_message = f"""Verify your email to activate your account.

Verification Link: {verification_url}

This link expires in 24 hours.

If you didn't request this email, please ignore it."""

    try:
        email = EmailMultiAlternatives(
            subject='Verify Your Email - AI Complaints System',
            body=plain_message,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[user.email],
        )
        
        # Only attach HTML if it was successfully rendered
        if html_message:
            email.attach_alternative(html_message, "text/html")
        
        # Send email with fail_silently=False to raise exceptions
        email.send(fail_silently=False)
        logger.info(f"✅ Verification email sent to {user.email}")
        
    except ImproperlyConfigured as e:
        logger.error(f"Email configuration error: {str(e)}")
        raise Exception("Email server is not properly configured. Please contact the system administrator.")
    except Exception as e:
        logger.error(f"Failed to send verification email to {user.email}: {str(e)}")
        raise Exception(f"Failed to send verification email: {str(e)}. Please try again later.")


def verify_email(request, token):
    """Handle email verification link click with proper error handling"""
    try:
        verification_token = EmailVerificationToken.objects.get(token=token)
    except EmailVerificationToken.DoesNotExist:
        messages.error(request, "Invalid verification link. Please check the link and try again, or sign up for a new account.")
        return render(request, 'accounts/verify_email_status.html', {'success': False})
    
    # Check if token is still valid
    if not verification_token.is_valid():
        messages.error(request, "Verification link has expired. Please sign up again to receive a new verification link.")
        try:
            user = verification_token.user
            verification_token.delete()
            user.delete()  # Delete expired user account
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error deleting expired user account: {str(e)}")
        return render(request, 'accounts/verify_email_status.html', {'success': False})
    
    # Mark user as active and email as verified
    try:
        user = verification_token.user
        user.is_active = True
        user.is_email_verified = True
        user.save()
        
        # Delete the verification token
        verification_token.delete()
        
        messages.success(
            request,
            "Email verified successfully! Your account is now active. You can now sign in with your credentials."
        )
        return render(request, 'accounts/verify_email_status.html', {'success': True})
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error verifying email for user: {str(e)}")
        messages.error(request, "An error occurred while verifying your email. Please contact support.")
        return render(request, 'accounts/verify_email_status.html', {'success': False})


def resend_verification_email(request):
    """Resend verification email to student who didn't receive it"""
    if request.method == "POST":
        email = request.POST.get("email")

        if not email:
            messages.error(request, "Please provide your email address.")
            return render(request, "accounts/resend_verification.html")

        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            messages.error(request, "No account found with this email address.")
            return render(request, "accounts/resend_verification.html")

        # Only resend if user is not verified
        if user.is_email_verified:
            messages.info(request, "This email is already verified. You can log in.")
            return redirect('login')

        try:
            # Delete any existing verification tokens
            EmailVerificationToken.objects.filter(user=user).delete()

            # Create new verification token
            verification_token = EmailVerificationToken.objects.create(user=user)

            # Send verification email
            send_verification_email(request, user, verification_token)

            messages.success(
                request,
                f"A new verification link has been sent to {email}. Please check your inbox (and spam folder) and click the link to verify your email."
            )
        except Exception as e:
            messages.error(request, f"Failed to resend verification email: {str(e)}")
            return render(request, "accounts/resend_verification.html")

        return redirect('login')

    return render(request, "accounts/resend_verification.html")


# ================= CHATBOT API =================

@require_http_methods(["POST"])
@login_required(login_url='login')
def chatbot_search(request):
    """
    API endpoint for chatbot to answer questions using Gemini AI
    Provides helpful responses about UDOM CIVE ICT and Academic issues
    """
    try:
        data = json.loads(request.body)
        query = data.get('query', '').strip()
        history = data.get('history', [])
        
        if not query or len(query) < 3:
            return JsonResponse({
                'success': False,
                'answer': 'Please provide a longer query (at least 3 characters).'
            })
        
        from .chatbot_service import get_chat_response
        result = get_chat_response(query, conversation_history=history)
        
        return JsonResponse({
            'success': True,
            'answer': result['answer'],
            'sources': result.get('sources', []),
            'from_gemini': result.get('from_gemini', False),
        })
            
    except Exception as e:
        return JsonResponse({
            'success': False,
            'answer': f'I encountered an error. Please try again or contact support.\n\nError: {str(e)}'
        })


def perform_google_search(query):
    """
    Perform a Google search and return relevant results
    Prioritizes UDOM (University of Dar es Salaam) results when applicable
    """
    try:
        # Add UDOM to query for UDOM-related searches
        if any(keyword in query.lower() for keyword in ['password', 'sr2', 'registration', 'academic', 'exam']):
            search_query = f"{query} site:udom.ac.tz OR site:dar.ac.tz"
        else:
            search_query = query
        
        # Build search URL
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        search_url = f"https://www.google.com/search?q={requests.utils.quote(search_query)}"
        
        # Try to get search results
        try:
            response = requests.get(search_url, headers=headers, timeout=5)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract search results
            sources = []
            answer_text = extract_answer_from_search(soup, query)
            
            # Extract top 3 result links
            for item in soup.select('div.g')[:3]:
                try:
                    title_elem = item.select_one('h3')
                    link_elem = item.select_one('a[href*="/url?q="]')
                    
                    if title_elem and link_elem:
                        title = title_elem.get_text()
                        link = link_elem.get('href', '')
                        # Clean up the URL
                        if '/url?q=' in link:
                            link = link.split('/url?q=')[1].split('&')[0]
                        
                        sources.append({'title': title, 'link': link})
                except:
                    continue
            
            return {
                'answer': answer_text,
                'sources': sources
            }
        
        except requests.RequestException:
            # Fallback: Generate a helpful response based on the query
            return generate_fallback_response(query)
    
    except Exception as e:
        return None


def extract_answer_from_search(soup, query):
    """Extract answer from Google search results"""
    try:
        # Try to get answer from knowledge panel or featured snippet
        knowledge_panel = soup.select_one('div[data-md*="answer"]')
        if knowledge_panel:
            return knowledge_panel.get_text()[:300]
        
        # Try featured snippet
        featured = soup.select_one('div.XcVN5e')
        if featured:
            return featured.get_text()[:300]
        
        # Try to get first search result snippet
        first_result = soup.select_one('div.s div.VwiC3b')
        if first_result:
            return first_result.get_text()[:300]
        
        # Fallback: construct answer from query
        return generate_fallback_response(query)['answer']
    except:
        return generate_fallback_response(query)['answer']


def generate_fallback_response(query):
    """Generate helpful fallback response based on query type"""
    query_lower = query.lower()
    
    # Check for specific keywords and provide relevant responses
    if 'password' in query_lower or 'reset' in query_lower:
        return {
            'answer': f'For password reset assistance, please contact the UDOM IT Support Center. You can reset your SR2 password by:\n\n1. Visit the UDOM portal (https://portal.udom.ac.tz)\n2. Click "Forgot Password"\n3. Enter your registration number\n4. Follow the email instructions to reset\n5. If issues persist, contact IT Help Desk: ithelpdesk@udom.ac.tz',
            'sources': []
        }
    elif 'registration' in query_lower or 'register' in query_lower:
        return {
            'answer': 'For course registration assistance, log into your student portal and follow these steps:\n\n1. Go to "Course Registration"\n2. Select your courses for the semester\n3. Check prerequisite requirements\n4. Submit your registration\n5. Confirm through your registered email\n\nFor support: registrar@udom.ac.tz',
            'sources': []
        }
    elif 'exam' in query_lower or 'mark' in query_lower or 'grade' in query_lower:
        return {
            'answer': 'For exam schedules, marks, and grades:\n\n1. Log into your student portal\n2. Check "Exam Schedule" section\n3. View your marks under "Academic Records"\n4. Contact your department if marks are missing\n5. Appeal process: Submit to Academic Affairs office\n\nFor inquiries: academics@udom.ac.tz',
            'sources': []
        }
    else:
        return {
            'answer': f'Based on your search for "{query}": I recommend checking the official UDOM website (www.udom.ac.tz) or contacting the appropriate department. For technical issues, contact IT Support.',
            'sources': []
        }


# ================= ADMIN FEATURES =================

@admin_required
def admin_users(request):
    """Admin page to manage all users"""
    users = User.objects.all().order_by('-date_joined')
    
    # Filter by role if provided
    role_filter = request.GET.get('role', '').strip()
    if role_filter and role_filter in ['student', 'staff', 'admin']:
        users = users.filter(role=role_filter)
    
    # Search functionality
    search_query = request.GET.get('search', '').strip()
    if search_query:
        users = users.filter(
            Q(username__icontains=search_query) |
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(email__icontains=search_query)
        )
    
    # Count users by role
    total_users = User.objects.count()
    student_count = User.objects.filter(role='student').count()
    staff_count = User.objects.filter(role='staff').count()
    admin_count = User.objects.filter(role='admin').count()
    
    context = {
        'users': users,
        'role_filter': role_filter,
        'search_query': search_query,
        'total_users': total_users,
        'student_count': student_count,
        'staff_count': staff_count,
        'admin_count': admin_count,
    }
    
    return render(request, 'accounts/admin_users.html', context)

@admin_required
def delete_user(request, user_id):
    """Delete a user account with proper cascade handling and atomicity"""
    import logging
    from django.db import transaction

    logger = logging.getLogger(__name__)
    
    try:
        user = User.objects.get(id=user_id)
        
        # Prevent self-deletion
        if user.id == request.user.id:
            messages.error(request, 'You cannot delete your own account!')
            return redirect('admin_users')
        
        username = user.username
        
        try:
            # Use atomic transaction to ensure all-or-nothing deletion
            with transaction.atomic():
                # Import needed models
                from complaints.models import Complaint, ComplaintResponse, AdminQueue
                
                # 1. First delete all complaint responses where this user is staff_member
                # These will cascade due to FK relationship but being explicit
                ComplaintResponse.objects.filter(staff_member=user).delete()
                
                # 2. Handle complaints assigned to this user - set to null before deletion
                # This prevents FK constraint violations
                Complaint.objects.filter(assigned_to=user).update(assigned_to=None)
                
                # 3. Delete EmailVerificationToken related to this user
                EmailVerificationToken.objects.filter(user=user).delete()
                
                # 4. Handle AdminQueue if it exists - set reviewed_by to null
                try:
                    AdminQueue.objects.filter(reviewed_by=user).update(reviewed_by=None)
                except Exception:
                    # AdminQueue table may not exist or field may not exist - that's ok
                    logger.warning("AdminQueue handling skipped - table may not exist")
                
                # 5. All Complaint objects where user is the creator will cascade delete
                # via Complaint.user FK with on_delete=models.CASCADE
                # But we do this manually to ensure proper error handling
                user_complaints = Complaint.objects.filter(user=user)
                complaint_ids = list(user_complaints.values_list('id', flat=True))
                if complaint_ids:
                    logger.info(f"Deleting {len(complaint_ids)} complaints from user {username}")
                    user_complaints.delete()
                
                # 6. Now safe to delete the user
                # All foreign key constraints have been handled
                user.delete()
                
                logger.info(f"✅ User successfully deleted: {username} (ID: {user_id})")
                messages.success(request, f'User "{username}" has been deleted successfully along with all their related data.')
                return redirect('admin_users')
        
        except IntegrityError as e:
            logger.error(f"Database integrity error while deleting user {username}: {str(e)}", exc_info=True)
            messages.error(request, f'Database error: User may have orphaned references. Please contact support.')
            return redirect('admin_users')
        except Exception as e:
            logger.error(f"Error during user deletion preparation for {username}: {str(e)}", exc_info=True)
            messages.error(request, f'Error deleting user: {str(e)}. Please try again or contact support.')
            return redirect('admin_users')
    
    except User.DoesNotExist:
        logger.warning(f"Attempted to delete non-existent user ID: {user_id}")
        messages.error(request, 'User not found!')
        return redirect('admin_users')
    except Exception as e:
        logger.error(f"Unexpected error in delete_user: {str(e)}", exc_info=True)
        messages.error(request, f'Unexpected error: {str(e)}')
        return redirect('admin_users')

@admin_required
def edit_user(request, user_id):
    """Edit user roles and details with proper validation"""
    import logging
    
    logger = logging.getLogger(__name__)
    
    try:
        user = User.objects.get(id=user_id)
        
        if request.method == 'POST':
            try:
                # Get form data
                first_name = request.POST.get('first_name', '').strip()
                last_name = request.POST.get('last_name', '').strip()
                email = request.POST.get('email', '').strip().lower()  # Normalize to lowercase
                role = request.POST.get('role', user.role).strip()
                department = request.POST.get('department', user.department or '').strip()
                registration_number = request.POST.get('registration_number', user.registration_number or '').strip()
                
                # Validate inputs
                errors = []
                
                # Email validation with case-insensitive check
                if email and email != user.email.lower():
                    if User.objects.filter(email__iexact=email).exclude(id=user.id).exists():
                        errors.append('Email is already in use by another user')
                
                # Role validation
                valid_roles = ['student', 'staff', 'admin']
                if role not in valid_roles:
                    errors.append(f'Invalid role: {role}')
                
                # Department validation for staff and admin
                if role in ['staff', 'admin'] and not department:
                    errors.append(f'Department is required for {role} users')
                
                # If there are errors, return form with errors
                if errors:
                    context = {
                        'user': user,
                        'role_choices': User.ROLE_CHOICES,
                        'error_list': errors,
                        'first_name': first_name,
                        'last_name': last_name,
                        'email': email,
                        'role': role,
                        'department': department,
                        'registration_number': registration_number,
                    }
                    return render(request, 'accounts/edit_user.html', context)
                
                # Update user
                user.first_name = first_name
                user.last_name = last_name
                user.email = email
                user.role = role
                user.department = department if department else None
                user.registration_number = registration_number if registration_number else None
                
                # If role changed from student to staff/admin, activate the user
                if user.role in ['staff', 'admin'] and not user.is_active:
                    user.is_active = True
                
                user.save()
                
                logger.info(f"✅ User {user.username} updated successfully by admin")
                messages.success(request, 'User updated successfully!')
                return redirect('admin_users')
            except Exception as e:
                logger.error(f"Error updating user {user.username}: {str(e)}", exc_info=True)
                messages.error(request, f'Error updating user: {str(e)}')
                context = {
                    'user': user,
                    'role_choices': User.ROLE_CHOICES,
                }
                return render(request, 'accounts/edit_user.html', context)
        
        context = {
            'user': user,
            'role_choices': User.ROLE_CHOICES,
        }
        return render(request, 'accounts/edit_user.html', context)
    except User.DoesNotExist:
        logger.warning(f"Attempted to edit non-existent user ID: {user_id}")
        messages.error(request, 'User not found!')
        return redirect('admin_users')
    except Exception as e:
        logger.error(f"Error in edit_user: {str(e)}", exc_info=True)
        messages.error(request, f'Error: {str(e)}')
        return redirect('admin_users')

@admin_required
def admin_analytics(request):
    """Admin analytics dashboard with charts and statistics"""
    all_complaints = Complaint.objects.all()
    
    # Complaint statistics
    total_complaints = all_complaints.count()
    pending_complaints = all_complaints.filter(status='Pending').count()
    in_progress_complaints = all_complaints.filter(status='In Progress').count()
    resolved_complaints = all_complaints.filter(status='Resolved').count()
    rejected_complaints = all_complaints.filter(status='Rejected').count()
    
    # Priority statistics
    urgent_complaints = all_complaints.filter(priority='Urgent').count()
    high_complaints = all_complaints.filter(priority='High').count()
    medium_complaints = all_complaints.filter(priority='Medium').count()
    low_complaints = all_complaints.filter(priority='Low').count()
    
    # User statistics
    total_users = User.objects.count()
    total_students = User.objects.filter(role='student').count()
    total_staff = User.objects.filter(role='staff').count()
    total_admins = User.objects.filter(role='admin').count()
    
    # Assigned complaints
    assigned_complaints = all_complaints.filter(assigned_to__isnull=False).count()
    unassigned_complaints = all_complaints.filter(assigned_to__isnull=True).count()
    
    # Most active staff
    staff_activity = User.objects.filter(role='staff').annotate(
        complaint_count=Count('assigned_complaints')
    ).order_by('-complaint_count')[:5]
    
    context = {
        'total_complaints': total_complaints,
        'pending_complaints': pending_complaints,
        'in_progress_complaints': in_progress_complaints,
        'resolved_complaints': resolved_complaints,
        'rejected_complaints': rejected_complaints,
        'urgent_complaints': urgent_complaints,
        'high_complaints': high_complaints,
        'medium_complaints': medium_complaints,
        'low_complaints': low_complaints,
        'total_users': total_users,
        'total_students': total_students,
        'total_staff': total_staff,
        'total_admins': total_admins,
        'assigned_complaints': assigned_complaints,
        'unassigned_complaints': unassigned_complaints,
        'staff_activity': staff_activity,
        
        # For JSON serialization in template
        'status_data': {
            'labels': ['Pending', 'In Progress', 'Resolved', 'Rejected'],
            'values': [pending_complaints, in_progress_complaints, resolved_complaints, rejected_complaints],
            'colors': ['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0']
        },
        'priority_data': {
            'labels': ['Urgent', 'High', 'Medium', 'Low'],
            'values': [urgent_complaints, high_complaints, medium_complaints, low_complaints],
            'colors': ['#FF4444', '#FF8844', '#FFBB44', '#44BB44']
        },
        'user_data': {
            'labels': ['Students', 'Staff', 'Admins'],
            'values': [total_students, total_staff, total_admins],
            'colors': ['#3498DB', '#2ECC71', '#E74C3C']
        }
    }
    
    return render(request, 'accounts/admin_analytics.html', context)

@admin_required
def admin_analytics_pdf(request):
    """Generate and download admin analytics report as PDF or Word document"""
    format_type = request.GET.get('format', 'pdf')  # Default to PDF
    
    all_complaints = Complaint.objects.all()
    
    # Complaint statistics
    total_complaints = all_complaints.count()
    pending_complaints = all_complaints.filter(status='Pending').count()
    in_progress_complaints = all_complaints.filter(status='In Progress').count()
    resolved_complaints = all_complaints.filter(status='Resolved').count()
    rejected_complaints = all_complaints.filter(status='Rejected').count()
    
    # Priority statistics
    urgent_complaints = all_complaints.filter(priority='Urgent').count()
    high_complaints = all_complaints.filter(priority='High').count()
    medium_complaints = all_complaints.filter(priority='Medium').count()
    low_complaints = all_complaints.filter(priority='Low').count()
    
    # User statistics
    total_users = User.objects.count()
    total_students = User.objects.filter(role='student').count()
    total_staff = User.objects.filter(role='staff').count()
    total_admins = User.objects.filter(role='admin').count()
    
    # Assigned complaints
    assigned_complaints = all_complaints.filter(assigned_to__isnull=False).count()
    unassigned_complaints = all_complaints.filter(assigned_to__isnull=True).count()
    
    # Most active staff
    staff_activity = User.objects.filter(role='staff').annotate(
        complaint_count=Count('assigned_complaints')
    ).order_by('-complaint_count')[:5]
    
    # Calculate resolution rate
    if total_complaints > 0:
        resolution_rate = ((resolved_complaints + rejected_complaints) / total_complaints * 100)
        pending_rate = (pending_complaints / total_complaints * 100)
    else:
        resolution_rate = 0
        pending_rate = 0
    
    if format_type == 'word':
        # Generate Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Admin Analytics Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('Complaint Management System Analytics')
        from datetime import datetime
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y %H:%M")}')
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph('This report provides a comprehensive overview of the complaint management system, including complaint statistics, user distribution, and staff activity metrics.')
        doc.add_paragraph()
        
        # Summary Statistics
        doc.add_heading('Summary Statistics', level=1)
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        summary_data = [
            ['Metric', 'Value'],
            ['Total Complaints', str(total_complaints)],
            ['Total Users', str(total_users)],
            ['Assigned Complaints', str(assigned_complaints)],
            ['Unassigned Complaints', str(unassigned_complaints)],
        ]
        
        for i, row_data in enumerate(summary_data):
            row = summary_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i == 0:
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Complaint Status Distribution
        doc.add_heading('Complaint Status Distribution', level=1)
        
        if total_complaints > 0:
            status_data = [
                ['Status', 'Count', 'Percentage'],
                ['Pending', str(pending_complaints), f"{(pending_complaints/total_complaints*100):.2f}%"],
                ['In Progress', str(in_progress_complaints), f"{(in_progress_complaints/total_complaints*100):.2f}%"],
                ['Resolved', str(resolved_complaints), f"{(resolved_complaints/total_complaints*100):.2f}%"],
                ['Rejected', str(rejected_complaints), f"{(rejected_complaints/total_complaints*100):.2f}%"],
            ]
        else:
            status_data = [
                ['Status', 'Count', 'Percentage'],
                ['Pending', str(pending_complaints), '0.00%'],
                ['In Progress', str(in_progress_complaints), '0.00%'],
                ['Resolved', str(resolved_complaints), '0.00%'],
                ['Rejected', str(rejected_complaints), '0.00%'],
            ]
        
        status_table = doc.add_table(rows=len(status_data), cols=3)
        status_table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(status_data):
            row = status_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i == 0:
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Complaint Priority Distribution
        doc.add_heading('Complaint Priority Distribution', level=1)
        
        if total_complaints > 0:
            priority_data = [
                ['Priority', 'Count', 'Percentage'],
                ['Urgent', str(urgent_complaints), f"{(urgent_complaints/total_complaints*100):.2f}%"],
                ['High', str(high_complaints), f"{(high_complaints/total_complaints*100):.2f}%"],
                ['Medium', str(medium_complaints), f"{(medium_complaints/total_complaints*100):.2f}%"],
                ['Low', str(low_complaints), f"{(low_complaints/total_complaints*100):.2f}%"],
            ]
        else:
            priority_data = [
                ['Priority', 'Count', 'Percentage'],
                ['Urgent', str(urgent_complaints), '0.00%'],
                ['High', str(high_complaints), '0.00%'],
                ['Medium', str(medium_complaints), '0.00%'],
                ['Low', str(low_complaints), '0.00%'],
            ]
        
        priority_table = doc.add_table(rows=len(priority_data), cols=3)
        priority_table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(priority_data):
            row = priority_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i == 0:
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # User Distribution
        doc.add_heading('User Distribution', level=1)
        
        if total_users > 0:
            user_data = [
                ['User Role', 'Count', 'Percentage'],
                ['Students', str(total_students), f"{(total_students/total_users*100):.2f}%"],
                ['Staff', str(total_staff), f"{(total_staff/total_users*100):.2f}%"],
                ['Admins', str(total_admins), f"{(total_admins/total_users*100):.2f}%"],
            ]
        else:
            user_data = [
                ['User Role', 'Count', 'Percentage'],
                ['Students', str(total_students), '0.00%'],
                ['Staff', str(total_staff), '0.00%'],
                ['Admins', str(total_admins), '0.00%'],
            ]
        
        user_table = doc.add_table(rows=len(user_data), cols=3)
        user_table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(user_data):
            row = user_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i == 0:
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Key Performance Metrics
        doc.add_heading('Key Performance Metrics', level=1)
        
        metrics_table = doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Light Grid Accent 1'
        metrics_data = [
            ['Metric', 'Value'],
            ['Resolution Rate', f"{resolution_rate:.1f}%"],
            ['Pending Rate', f"{pending_rate:.1f}%"],
            ['Unassigned Complaints', str(unassigned_complaints)],
            ['Critical Issues (Urgent)', str(urgent_complaints)],
        ]
        
        for i, row_data in enumerate(metrics_data):
            row = metrics_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i == 0:
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Top Staff Members
        if staff_activity:
            doc.add_heading('Top Staff Members (by Assigned Complaints)', level=1)
            
            staff_data = [['Staff Name', 'Assigned Complaints']]
            for staff in staff_activity:
                staff_data.append([staff.get_full_name or staff.username, str(staff.complaint_count)])
            
            staff_table = doc.add_table(rows=len(staff_data), cols=2)
            staff_table.style = 'Light Grid Accent 1'
            
            for i, row_data in enumerate(staff_data):
                row = staff_table.rows[i]
                for j, cell_data in enumerate(row_data):
                    row.cells[j].text = cell_data
                    if i == 0:
                        row.cells[j].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
        
        # Footer
        doc.add_paragraph('AI Complaints Management System - University of Dodoma')
        doc.add_paragraph(f'Report generated on {datetime.now().strftime("%B %d, %Y %H:%M")}')
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="admin_analytics_report.docx"'
        
        return response
    else:
        # Generate PDF using ReportLab
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=18,
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=12
        )
        
        # Title
        story.append(Paragraph("Admin Analytics Report", title_style))
        story.append(Paragraph("Complaint Management System Analytics", styles['Normal']))
        from datetime import datetime
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 0.2 * inch))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        story.append(Paragraph(
            f"This report provides a comprehensive overview of the complaint management system, including complaint statistics, user distribution, and staff activity metrics.",
            styles['Normal']
        ))
        story.append(Spacer(1, 0.2 * inch))
        
        # Summary Statistics
        story.append(Paragraph("Summary Statistics", heading_style))
        
        summary_data = [
            ['Metric', 'Value'],
            ['Total Complaints', str(total_complaints)],
            ['Total Users', str(total_users)],
            ['Assigned Complaints', str(assigned_complaints)],
            ['Unassigned Complaints', str(unassigned_complaints)],
        ]
        
        summary_table = Table(summary_data, colWidths=[3 * inch, 2 * inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#1e3a8a')),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.3 * inch))
        
        # Complaint Status Distribution
        story.append(Paragraph("Complaint Status Distribution", heading_style))
        
        if total_complaints > 0:
            status_data = [
                ['Status', 'Count', 'Percentage'],
                ['Pending', str(pending_complaints), f"{(pending_complaints/total_complaints*100):.2f}%"],
                ['In Progress', str(in_progress_complaints), f"{(in_progress_complaints/total_complaints*100):.2f}%"],
                ['Resolved', str(resolved_complaints), f"{(resolved_complaints/total_complaints*100):.2f}%"],
                ['Rejected', str(rejected_complaints), f"{(rejected_complaints/total_complaints*100):.2f}%"],
            ]
        else:
            status_data = [
                ['Status', 'Count', 'Percentage'],
                ['Pending', str(pending_complaints), '0.00%'],
                ['In Progress', str(in_progress_complaints), '0.00%'],
                ['Resolved', str(resolved_complaints), '0.00%'],
                ['Rejected', str(rejected_complaints), '0.00%'],
            ]
        
        status_table = Table(status_data, colWidths=[2 * inch, 1.5 * inch, 1.5 * inch])
        status_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (2, 0), colors.HexColor('#1e3a8a')),
            ('TEXTCOLOR', (0, 0), (2, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(status_table)
        story.append(Spacer(1, 0.3 * inch))
        
        # Complaint Priority Distribution
        story.append(Paragraph("Complaint Priority Distribution", heading_style))
        
        if total_complaints > 0:
            priority_data = [
                ['Priority', 'Count', 'Percentage'],
                ['Urgent', str(urgent_complaints), f"{(urgent_complaints/total_complaints*100):.2f}%"],
                ['High', str(high_complaints), f"{(high_complaints/total_complaints*100):.2f}%"],
                ['Medium', str(medium_complaints), f"{(medium_complaints/total_complaints*100):.2f}%"],
                ['Low', str(low_complaints), f"{(low_complaints/total_complaints*100):.2f}%"],
            ]
        else:
            priority_data = [
                ['Priority', 'Count', 'Percentage'],
                ['Urgent', str(urgent_complaints), '0.00%'],
                ['High', str(high_complaints), '0.00%'],
                ['Medium', str(medium_complaints), '0.00%'],
                ['Low', str(low_complaints), '0.00%'],
            ]
        
        priority_table = Table(priority_data, colWidths=[2 * inch, 1.5 * inch, 1.5 * inch])
        priority_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (2, 0), colors.HexColor('#dc3545')),
            ('TEXTCOLOR', (0, 0), (2, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(priority_table)
        story.append(Spacer(1, 0.3 * inch))
        
        # User Distribution
        story.append(Paragraph("User Distribution", heading_style))
        
        if total_users > 0:
            user_data = [
                ['User Role', 'Count', 'Percentage'],
                ['Students', str(total_students), f"{(total_students/total_users*100):.2f}%"],
                ['Staff', str(total_staff), f"{(total_staff/total_users*100):.2f}%"],
                ['Admins', str(total_admins), f"{(total_admins/total_users*100):.2f}%"],
            ]
        else:
            user_data = [
                ['User Role', 'Count', 'Percentage'],
                ['Students', str(total_students), '0.00%'],
                ['Staff', str(total_staff), '0.00%'],
                ['Admins', str(total_admins), '0.00%'],
            ]
        
        user_table = Table(user_data, colWidths=[2 * inch, 1.5 * inch, 1.5 * inch])
        user_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (2, 0), colors.HexColor('#1e3a8a')),
            ('TEXTCOLOR', (0, 0), (2, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(user_table)
        story.append(Spacer(1, 0.3 * inch))
        
        # Key Performance Metrics
        story.append(Paragraph("Key Performance Metrics", heading_style))
        
        metrics_data = [
            ['Metric', 'Value'],
            ['Resolution Rate', f"{resolution_rate:.1f}%"],
            ['Pending Rate', f"{pending_rate:.1f}%"],
            ['Unassigned Complaints', str(unassigned_complaints)],
            ['Critical Issues (Urgent)', str(urgent_complaints)],
        ]
        
        metrics_table = Table(metrics_data, colWidths=[3 * inch, 2 * inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#1e3a8a')),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(metrics_table)
        story.append(Spacer(1, 0.3 * inch))
        
        # Top Staff Members
        if staff_activity:
            story.append(Paragraph("Top Staff Members (by Assigned Complaints)", heading_style))
            
            staff_data = [['Staff Name', 'Assigned Complaints']]
            for staff in staff_activity:
                staff_data.append([staff.get_full_name or staff.username, str(staff.complaint_count)])
            
            staff_table = Table(staff_data, colWidths=[3.5 * inch, 1.5 * inch])
            staff_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#1e3a8a')),
                ('TEXTCOLOR', (0, 0), (1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(staff_table)
            story.append(Spacer(1, 0.3 * inch))
        
        # Footer
        story.append(Paragraph("AI Complaints Management System - University of Dodoma", styles['Normal']))
        story.append(Paragraph(f"Report generated on {datetime.now().strftime('%B %d, %Y %H:%M')}", styles['Normal']))
        
        # Build PDF
        doc.build(story)
        
        # Create response
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="admin_analytics_report.pdf"'
        
        return response
